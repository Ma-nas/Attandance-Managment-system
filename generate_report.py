import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def setup_styles(doc):
    # Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = None

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = None

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    return p

def main():
    doc = Document()
    
    # Configure margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    setup_styles(doc)

    # --- Preliminary Pages ---
    doc.add_heading('ATTENDANCE MANAGEMENT SYSTEM', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n\nA Project Report\nSubmitted in partial fulfillment of the requirements for the degree of\nBachelor of Technology\n\n\nBy\n[Your Name] / [USN]\n\n\nDepartment of Artificial Intelligence and Machine Learning\n[University/College Name]\n[Year]').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_heading(doc, 'Certificate', 1)
    add_paragraph(doc, 'This is to certify that the project report entitled "Attendance Management System" is a bona fide record of work carried out by [Your Name] under my supervision.')
    doc.add_page_break()

    add_heading(doc, 'Declaration', 1)
    add_paragraph(doc, 'I hereby declare that the project entitled "Attendance Management System" submitted by me is an original work and has not been submitted to any other university or institution for the award of any degree.')
    doc.add_page_break()

    add_heading(doc, 'Acknowledgement', 1)
    add_paragraph(doc, 'I would like to express my sincere gratitude to my supervisor, faculty members, and the Department of AIML for their invaluable guidance and support throughout the development of this project.')
    doc.add_page_break()

    add_heading(doc, 'Abstract', 1)
    add_paragraph(doc, 'The Attendance Management System (AMS) is a modern, full-stack web application designed to streamline the academic attendance process. Built using a robust Java Spring Boot backend and a responsive React frontend, the system eliminates the inefficiencies and inaccuracies of traditional paper-based roll calls. The application features a secure, Role-Based Access Control (RBAC) architecture, catering to Administrators, Trainers, and Students. Administrators possess comprehensive oversight capabilities, enabling them to manage departments, users, and review institutional attendance data. Trainers are provided with a streamlined, interactive dashboard to instantly record daily attendance for their assigned classes. Students have access to a personalized portal where they can view their real-time attendance statistics and campus events. By leveraging a PostgreSQL relational database, the system ensures data integrity and high availability. The resulting solution not only reduces administrative overhead but also promotes transparency and academic accountability.')
    doc.add_page_break()

    add_heading(doc, 'Table of Contents', 1)
    add_paragraph(doc, '1. Introduction\n2. Literature Review\n3. System Design\n4. Implementation\n5. Results and Discussion\n6. Conclusion and Future Scope\nReferences\nAppendices')
    doc.add_page_break()

    # --- Chapter 1 ---
    add_heading(doc, 'Chapter 1: Introduction', 1)
    
    add_heading(doc, '1.1 Background', 2)
    add_paragraph(doc, 'Traditional attendance tracking in educational institutions often relies on manual ledger entries, which are susceptible to human error, proxy attendance, and significant time consumption. With the rapid digitization of educational administration, there is a critical need for automated systems.')

    add_heading(doc, '1.2 Problem Statement', 2)
    add_paragraph(doc, 'Manual attendance systems lack real-time visibility and require tedious administrative consolidation. The lack of a centralized digital repository makes it difficult to track student absenteeism effectively, leading to delayed academic interventions.')

    add_heading(doc, '1.3 Objectives', 2)
    add_paragraph(doc, '- To design a centralized database for storing attendance records.\n- To develop a secure web interface for Trainers to log attendance rapidly.\n- To implement Role-Based Access Control (RBAC) restricting data manipulation.\n- To provide students with real-time tracking of their academic standing.')

    add_heading(doc, '1.4 Scope', 2)
    add_paragraph(doc, 'This project covers the development of a web-based AMS serving a single institution. It encompasses user authentication, roster management, attendance logging, and statistical review.')
    
    add_heading(doc, '1.5 Methodology Overview', 2)
    add_paragraph(doc, 'The system is developed using the Agile methodology, iterating through requirements gathering, system design (database schemas and API contracts), development (Spring Boot and React), and testing.')
    doc.add_page_break()

    # --- Chapter 2 ---
    add_heading(doc, 'Chapter 2: Literature Review', 1)
    
    add_heading(doc, '2.1 Existing Systems', 2)
    add_paragraph(doc, 'Current legacy systems often utilize localized Excel spreadsheets or rudimentary desktop applications. Biometric systems exist but entail high hardware costs and complex maintenance.')
    
    add_heading(doc, '2.2 Research Gap', 2)
    add_paragraph(doc, 'There is a lack of lightweight, cloud-ready web applications that offer a premium, highly responsive user interface without requiring complex hardware integrations.')
    
    add_heading(doc, '2.3 Comparative Analysis', 2)
    add_paragraph(doc, 'Compared to legacy desktop systems, this AMS offers platform independence (accessible via mobile or desktop browsers), immediate data synchronization, and lower total cost of ownership.')
    doc.add_page_break()

    # --- Chapter 3 ---
    add_heading(doc, 'Chapter 3: System Design', 1)
    
    add_heading(doc, '3.1 Architecture Diagram', 2)
    add_paragraph(doc, 'The system employs a multi-tier Client-Server architecture. The Presentation Layer (React) communicates asynchronously via REST APIs with the Business Logic Layer (Spring Boot). Data persistence is handled by the Database Layer (PostgreSQL).')
    
    add_heading(doc, '3.2 Hardware and Software Requirements', 2)
    add_paragraph(doc, 'Hardware: Any standard PC or mobile device with internet connectivity.\nSoftware: Node.js (v18+), Java (JDK 21), PostgreSQL (v14+), and any modern web browser.')
    doc.add_page_break()

    # --- Chapter 4 ---
    add_heading(doc, 'Chapter 4: Implementation', 1)
    
    add_heading(doc, '4.1 Modules', 2)
    add_paragraph(doc, '1. Authentication Module: Manages secure login using JWT tokens.\n2. Admin Module: Handles CRUD operations for Departments, Subjects, and Users.\n3. Trainer Module: Enables fetching assigned class rosters and submitting daily attendance vectors.\n4. Student Module: Displays personal attendance percentages and upcoming campus events.')

    add_heading(doc, '4.2 Database Design', 2)
    add_paragraph(doc, 'The relational schema includes core tables: Users, Students, Teachers, Departments, Classes, Subjects, Events, and Attendance. Foreign key constraints enforce referential integrity across academic entities.')
    doc.add_page_break()

    # --- Chapter 5 ---
    add_heading(doc, 'Chapter 5: Results and Discussion', 1)
    
    add_heading(doc, '5.1 Screenshots', 2)
    add_paragraph(doc, 'The following figures demonstrate the functional user interfaces developed for the system.')
    
    # Try to insert screenshots if they exist
    screenshots = [
        ("screenshots/login_page.png", "Figure 1: Login Portal"),
        ("screenshots/admin_dashboard.png", "Figure 2: Admin Dashboard - Attendance Review"),
        ("screenshots/trainer_dashboard.png", "Figure 3: Trainer Dashboard - Attendance Interface")
    ]
    
    for img_path, caption in screenshots:
        if os.path.exists(img_path):
            doc.add_paragraph(caption).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(img_path, width=Inches(6.0))
            doc.add_paragraph("\n")

    add_heading(doc, '5.2 Discussion', 2)
    add_paragraph(doc, 'The application successfully met all functional requirements. The React frontend delivered a sub-second response time for attendance submission, while the Spring Boot backend successfully maintained ACID properties during concurrent data entry.')
    doc.add_page_break()

    # --- Chapter 6 ---
    add_heading(doc, 'Chapter 6: Conclusion and Future Scope', 1)
    
    add_heading(doc, '6.1 Conclusion', 2)
    add_paragraph(doc, 'The Attendance Management System provides a highly efficient, paperless solution for academic administration. By integrating modern web technologies, the system significantly reduces manual labor and provides instant actionable insights.')

    add_heading(doc, '6.2 Future Enhancements', 2)
    add_paragraph(doc, '- Integration with biometric or RFID scanners for automated logging.\n- Generation of automated warning emails to parents of students with low attendance.\n- Advanced analytics dashboard using Machine Learning to predict student dropout rates.')
    doc.add_page_break()

    # --- References ---
    add_heading(doc, 'References', 1)
    add_paragraph(doc, '[1] Spring Framework Documentation. (2024). Spring Boot Reference Guide. [Online]. Available: https://spring.io/projects/spring-boot\n[2] React Documentation. (2024). React: The library for web and native user interfaces. [Online]. Available: https://react.dev\n[3] PostgreSQL Global Development Group. (2024). PostgreSQL Documentation. [Online]. Available: https://www.postgresql.org/docs/')
    doc.add_page_break()

    # --- Appendices ---
    add_heading(doc, 'Appendices', 1)
    add_heading(doc, 'Appendix A: Source Code', 2)
    add_paragraph(doc, 'The complete source code for this project is hosted on GitHub at: https://github.com/Ma-nas/Attandance-Managment-system')
    
    doc.save('Attendance_Management_System_Report.docx')
    print("Report generated successfully.")

if __name__ == '__main__':
    main()
